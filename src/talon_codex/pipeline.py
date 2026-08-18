"""High-level orchestration. Nothing in this module runs on import."""

from __future__ import annotations

import json
import platform
from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd
import torch
import joblib

from .analysis.comparisons import compare_suite
from .config import ResearchConfig
from .data import audit_metadata, build_cross_validation_bundle, build_data_bundle, build_external_loader, build_grouped_cv, export_cleaned_binary_masks, load_metadata
from .evaluation import analyze_and_export_segmentation, case_aggregation, collect_raw_predictions, evaluate_external_dataset, evaluate_model, generate_gradcam_artifacts, write_artifact_manifest
from .models import MaskGuidedMultiTaskUNet, baseline_width_from_checkpoint, build_talon, build_training_spatial_prior
from .reporting import render_evaluation_figures, render_training_history, write_q1_workbook
from .training import train_experiment
from .analysis.statistics import classification_metrics, select_segmentation_threshold


def prepare_data(config: ResearchConfig, seed: int | None = None, hash_audit: bool = True) -> dict[str, Any]:
    seed = config.seed if seed is None else int(seed)
    bundle = build_data_bundle(config, seed)
    audit_dir = config.output_root / "audit"
    duplicate_dir = audit_dir / "duplicate_checks"
    mask_dir = audit_dir / "mask_checks"
    split_dir = audit_dir / "split_assignments"
    for directory in (duplicate_dir, mask_dir, split_dir):
        directory.mkdir(parents=True, exist_ok=True)
    audits = audit_metadata(pd.concat(bundle.frames.values(), ignore_index=True), config, include_hashes=hash_audit)
    audits["mask_audit"].to_csv(mask_dir / "mask_audit.csv", index=False)
    audits["pixel_hashes"].to_csv(duplicate_dir / "pixel_hashes.csv", index=False)
    audits["duplicate_images"].to_csv(duplicate_dir / "duplicate_images.csv", index=False)
    audits["duplicate_rows"].to_csv(duplicate_dir / "duplicate_metadata_rows.csv", index=False)
    audits["duplicate_paths"].to_csv(duplicate_dir / "duplicate_resolved_paths.csv", index=False)
    audits["duplicate_file_bytes"].to_csv(duplicate_dir / "duplicate_file_bytes.csv", index=False)
    duplicate_table = audits["duplicate_images"]
    if not duplicate_table.empty and (duplicate_table["DuplicateRisk"] == "cross_split_leakage").any():
        raise ValueError("Pixel-identical images cross the locked subject split; see audit/duplicate_checks/duplicate_images.csv")
    if bool(audits["mask_audit"]["IsEmptyAfterThreshold"].any()):
        raise ValueError("At least one lesion mask is empty after gray >127 binarization; see audit/mask_checks/mask_audit.csv")
    if not bool(audits["mask_audit"]["ImageMaskShapeMatch"].all()):
        raise ValueError("At least one image/mask dimension mismatch exists; see audit/mask_checks/mask_audit.csv")
    cleaned_root = config.output_root / "cleaned_masks"
    cleaned_manifest_path = cleaned_root / "cleaned_mask_manifest.csv"
    cleaned_mask_manifest = pd.read_csv(cleaned_manifest_path) if cleaned_manifest_path.exists() else export_cleaned_binary_masks(pd.concat(bundle.frames.values(), ignore_index=True), config, cleaned_root)
    assignments = pd.concat([
        frame[["SubjectID", "CaseID", "MaskedPatientName", "ClassName", "Split"]].drop_duplicates()
        for frame in bundle.frames.values()
    ], ignore_index=True).sort_values(["Split", "SubjectID", "CaseID"])
    assignments.to_csv(split_dir / f"split_assignments_seed_{seed}.csv", index=False)
    counts = pd.concat(bundle.frames.values()).groupby(["Split", "ClassName"], observed=True).agg(subjects=("SubjectID", "nunique"), metadata_cases=("CaseID", "nunique"), slices=("ImageName", "size")).reset_index()
    counts.to_csv(split_dir / f"split_counts_seed_{seed}.csv", index=False)
    (split_dir / f"roi_edges_train_only_seed_{seed}.json").write_text(json.dumps(bundle.roi_edges), encoding="utf-8")
    config.snapshot(audit_dir / "resolved_config.json", {
        "python": platform.python_version(), "platform": platform.platform(), "torch": torch.__version__, "numpy": np.__version__, "pandas": pd.__version__,
    })
    return {"bundle": bundle, "audits": audits, "assignments": assignments, "counts": counts, "cleaned_mask_manifest": cleaned_mask_manifest}


def export_cross_validation_plan(config: ResearchConfig) -> pd.DataFrame:
    locked_bundle = build_data_bundle(config, config.seed)
    frame = pd.concat([locked_bundle.frames["train"], locked_bundle.frames["validation"]], ignore_index=True)
    rows = []
    for fold, (train_indices, validation_indices) in enumerate(build_grouped_cv(frame, config), start=1):
        for role, indices in (("train", train_indices), ("validation", validation_indices)):
            part = frame.iloc[indices][["SubjectID", "CaseID", "MaskedPatientName", "ClassName"]].drop_duplicates()
            part.insert(0, "role", role)
            part.insert(0, "fold", fold)
            rows.append(part)
    result = pd.concat(rows, ignore_index=True)
    destination = config.output_root / "cross_validation" / "grouped_cross_validation_plan.csv"
    destination.parent.mkdir(parents=True, exist_ok=True)
    result.to_csv(destination, index=False)
    return result


def run_one_experiment(config: ResearchConfig, experiment: str, seed: int) -> dict[str, Any]:
    prepared = prepare_data(config, seed, hash_audit=False)
    bundle = prepared["bundle"]
    model, history, checkpoint = train_experiment(config, bundle, experiment, seed)
    run_dir = config.run_dir(experiment, seed)
    render_training_history(history, run_dir / "figures" / "training")
    evaluation = evaluate_model(model, bundle.loaders, config, experiment, seed)
    device = next(model.parameters()).device
    evaluation_dir = Path(evaluation["output_dir"])
    xai = generate_gradcam_artifacts(model, bundle.loaders["test"], evaluation_dir / "xai", device)
    artifact_manifest = write_artifact_manifest(evaluation_dir, config, experiment, seed)
    return {"checkpoint": checkpoint, "history": history, "evaluation": evaluation, "xai": xai, "artifact_manifest": artifact_manifest}


def run_experiment_suite(config: ResearchConfig, include_repeated_seeds: bool = True, include_ablations: bool = True) -> None:
    experiments = list(config.section("experiments")["primary"])
    if include_ablations:
        experiments.extend(config.section("experiments")["ablations"])
    seeds = list(config.get("repeated_seeds")) if include_repeated_seeds else [config.seed]
    for seed in seeds:
        for experiment in experiments:
            run_one_experiment(config, str(experiment), int(seed))


def run_statistical_comparisons(config: ResearchConfig, include_ablations: bool = True) -> pd.DataFrame:
    experiments = list(config.section("experiments")["primary"])
    if include_ablations:
        experiments.extend(config.section("experiments")["ablations"])
    table = compare_suite(config.output_root, experiments, list(config.get("repeated_seeds")), int(config.section("statistics")["bootstrap_iterations"]))
    destination = config.output_root / "statistical_comparisons"
    destination.mkdir(parents=True, exist_ok=True)
    table.to_csv(destination / "paired_model_comparisons_holm.csv", index=False)
    seed_summary = table.groupby(["comparison", "endpoint"], observed=True).agg(seeds=("seed", "nunique"), mean_effect=("effect", "mean"), sd_effect=("effect", "std"), median_p_holm=("p_holm", "median")).reset_index()
    write_q1_workbook({"paired_tests": table, "seed_summary": seed_summary}, destination / "q1_model_comparison_tables.xlsx")
    return table


def compare_teacher_checkpoints(config: ResearchConfig, experiment: str, seed: int) -> pd.DataFrame:
    """Compare Teacher 1–5 on validation only; the locked test is never used for phase selection."""
    bundle = build_data_bundle(config, seed)
    run_dir = config.run_dir(experiment, seed)
    output_root = run_dir / "teacher_checkpoint_comparison"
    prior = build_training_spatial_prior([Path(path) for path in bundle.frames["train"]["ResolvedRoiPath"]], int(config.get("mask_threshold", 127)))
    rows = []
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    for checkpoint in sorted((run_dir / "checkpoints").glob("phase_*_best.pt")):
        model = build_talon(config.section("model"), prior, experiment).to(device)
        model.load_state_dict(torch.load(checkpoint, map_location=device, weights_only=False)["model_state_dict"])
        raw = collect_raw_predictions(model, bundle.loaders["validation"], device)
        threshold, _ = select_segmentation_threshold(raw, config.section("selection")["segmentation_thresholds"])
        phase_dir = output_root / checkpoint.stem
        slice_metrics, pred_components, _ = analyze_and_export_segmentation(raw, threshold, phase_dir, config.section("components"))
        cases = case_aggregation(raw, str(config.section("selection")["case_aggregation_primary"]), int(config.section("selection")["top_k"]))
        cls = classification_metrics(cases["class_id"], cases["class_probability"], float(config.section("selection")["classification_threshold"]), int(config.section("statistics")["ece_bins"]))
        xai = generate_gradcam_artifacts(model, bundle.loaders["validation"], phase_dir / "xai", device)
        rows.append({
            "checkpoint": checkpoint.name, "segmentation_threshold": threshold,
            "mean_dice": slice_metrics["dice"].mean(), "mean_iou": slice_metrics["iou"].mean(),
            "off_target_fp_occurrences": int((pred_components["category"] == "off_target_fp").sum()),
            "poorly_matched_occurrences": int((pred_components["category"] == "poorly_matched").sum()),
            "xai_energy_inside_gt": xai["xai_energy_inside_gt"].mean(),
            "xai_off_target_attention": xai["xai_off_target_attention"].mean(), **cls,
        })
    table = pd.DataFrame(rows)
    output_root.mkdir(parents=True, exist_ok=True)
    table.to_csv(output_root / "teacher_checkpoint_validation_comparison.csv", index=False)
    return table


def run_grouped_cross_validation(config: ResearchConfig, experiment: str = "TALON_FULL") -> pd.DataFrame:
    """Train/evaluate all subject-grouped development folds without the locked test."""
    locked_bundle = build_data_bundle(config, config.seed)
    frame = pd.concat([locked_bundle.frames["train"], locked_bundle.frames["validation"]], ignore_index=True)
    summaries = []
    for fold, (train_indices, validation_indices) in enumerate(build_grouped_cv(frame, config), start=1):
        fold_root = config.output_root / "cross_validation" / experiment / f"fold_{fold:02d}"
        fold_raw = dict(config.raw); fold_raw["output_root"] = str(fold_root)
        fold_config = ResearchConfig(fold_raw, config.source_path)
        fold_seed = config.seed + fold * 1000
        bundle = build_cross_validation_bundle(frame, train_indices, validation_indices, fold_config, fold_seed)
        model, _, checkpoint = train_experiment(fold_config, bundle, experiment, fold_seed)
        device = next(model.parameters()).device
        raw = collect_raw_predictions(model, bundle.loaders["validation"], device)
        threshold, _ = select_segmentation_threshold(raw, config.section("selection")["segmentation_thresholds"])
        cases = case_aggregation(raw, str(config.section("selection")["case_aggregation_primary"]), int(config.section("selection")["top_k"]))
        cls = classification_metrics(cases["class_id"], cases["class_probability"], float(config.section("selection")["classification_threshold"]), int(config.section("statistics")["ece_bins"]))
        dice = np.mean([((2 * ((prob >= threshold) & target).sum() + 1e-6) / ((prob >= threshold).sum() + target.sum() + 1e-6)) for prob, target in zip(raw["seg_probability_map"], raw["target_mask"])])
        summaries.append({"fold": fold, "seed": fold_seed, "checkpoint": str(checkpoint), "validation_dice": dice, **cls})
    result = pd.DataFrame(summaries)
    destination = config.output_root / "cross_validation" / experiment
    destination.mkdir(parents=True, exist_ok=True)
    result.to_csv(destination / "grouped_cv_results.csv", index=False)
    return result


def extract_clinical_source_inventory(config: ResearchConfig) -> pd.DataFrame:
    """Extract thesis DOCX evidence candidates without promoting them to verified facts."""
    from docx import Document
    source = config.resolve_path(config.get("clinical_document_path"))
    document = Document(source)
    blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            blocks.append(f"TABLE {table_index} ROW {row_index}: " + " | ".join(cell.text.strip() for cell in row.cells))
    output = config.output_root / "reports" / "clinical_source_inventory"
    output.mkdir(parents=True, exist_ok=True)
    (output / "igm_tez_metni_extracted.txt").write_text("\n".join(blocks), encoding="utf-8")
    fields = {
        "ethics": ["etik", "kurul", "onay"], "consent": ["onam", "muaf"],
        "retrospective_design": ["retrospektif", "prospektif"], "center": ["hastane", "merkez"],
        "study_dates": ["tarih", "yıl", "arasında"], "diagnostic_standard": ["tanı", "histopat"],
        "inclusion_exclusion": ["dahil", "dışla", "hariç", "kriter"],
        "participant_flow_107_to_103": ["107", "103", "akış", "dışlanan"],
        "demographics": ["yaş", "demograf", "cinsiyet"],
        "mri_hardware": ["üretici", "model", "tesla", "coil", "bobin"],
        "mri_acquisition": ["tr/te", "flip", "fov", "matris", "kesit", "kontrast", "doz", "faz", "seri"],
        "annotation_personnel": ["radyolog", "deneyim", "kör", "uzlaşı", "gözlemci"],
        "annotation_format_conversion": ["segment", "anotasyon", "maske", "format", "dönüş"],
        "node_radiology_pathology_matching": ["aksiller", "lenf", "hedef nod", "patolojik nod", "birebir", "eşleştir"],
    }
    records = []
    for field, keywords in fields.items():
        candidates = [text for text in blocks if any(keyword.casefold() in text.casefold() for keyword in keywords)]
        records.append({"field": field, "status": "candidate_text_requires_hospital_verification" if candidates else "not_found_pending", "candidate_count": len(candidates), "candidate_text": " || ".join(candidates[:10])})
    inventory = pd.DataFrame(records)
    inventory.to_csv(output / "clinical_field_inventory.csv", index=False)
    return inventory


def run_external_validation(config: ResearchConfig, metadata_path: Path, experiment: str, seed: int) -> dict[str, Any]:
    """Evaluate a separately supplied cohort with frozen weights and thresholds."""
    bundle = build_data_bundle(config, seed)
    run_dir = config.run_dir(experiment, seed)
    checkpoint_path = run_dir / "checkpoints" / "best_overall.pt"
    if not checkpoint_path.exists():
        raise FileNotFoundError(checkpoint_path)
    prior = build_training_spatial_prior([Path(path) for path in bundle.frames["train"]["ResolvedRoiPath"]], int(config.get("mask_threshold", 127)))
    checkpoint = torch.load(checkpoint_path, map_location="cpu", weights_only=False)
    if experiment == "UNET_MTL_MASK_GUIDED":
        model = MaskGuidedMultiTaskUNet(baseline_width_from_checkpoint(checkpoint, config.section("model")), float(config.section("model")["classifier_dropout"]))
    else:
        model = build_talon(config.section("model"), prior, experiment)
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    model.load_state_dict(checkpoint["model_state_dict"])
    model.to(device)
    _, external_loader = build_external_loader(metadata_path, config, bundle.roi_edges)
    thresholds = json.loads((run_dir / "evaluation" / "reproducibility" / "locked_thresholds.json").read_text(encoding="utf-8"))
    calibrator_path = run_dir / "evaluation" / "reproducibility" / "slice_calibrator.joblib"
    calibrator = joblib.load(calibrator_path) if calibrator_path.exists() else None
    case_calibrator_path = run_dir / "evaluation" / "reproducibility" / "case_calibrator.joblib"
    case_calibrator = joblib.load(case_calibrator_path) if case_calibrator_path.exists() else None
    output = config.output_root / "external_validation" / experiment / f"seed_{seed}"
    output.mkdir(parents=True, exist_ok=True)
    results = evaluate_external_dataset(model, external_loader, float(thresholds["segmentation_threshold"]), float(thresholds["classification_threshold"]), config, output, device, calibrator, case_calibrator, thresholds.get("calibration_method"))
    results["slice_metrics"].to_csv(output / "external_slice_segmentation_metrics.csv", index=False)
    results["case_rows"].to_csv(output / "external_case_classification_predictions.csv", index=False)
    results["predicted_components"].to_csv(output / "external_predicted_components.csv", index=False)
    results["ground_truth_components"].to_csv(output / "external_ground_truth_components.csv", index=False)
    (output / "external_classification_metrics.json").write_text(json.dumps(results["classification"], indent=2), encoding="utf-8")
    if results["classification_calibrated"] is not None:
        (output / "external_classification_metrics_calibrated.json").write_text(json.dumps(results["classification_calibrated"], indent=2), encoding="utf-8")
        results["calibrated_case_rows"].to_csv(output / "external_case_classification_predictions_calibrated.csv", index=False)
    xai = generate_gradcam_artifacts(model, external_loader, output / "xai", device)
    render_evaluation_figures(results["case_rows"], output / "figures" / "raw", float(thresholds["classification_threshold"]), int(config.section("statistics")["ece_bins"]), level="external_case")
    workbook_tables = {"case_predictions": results["case_rows"], "slice_segmentation": results["slice_metrics"], "pred_components": results["predicted_components"], "gt_components": results["ground_truth_components"], "xai": xai}
    if results["calibrated_case_rows"] is not None:
        render_evaluation_figures(results["calibrated_case_rows"], output / "figures" / "calibrated", float(thresholds["classification_threshold"]), int(config.section("statistics")["ece_bins"]), level="external_case_calibrated")
        workbook_tables["case_predictions_calibrated"] = results["calibrated_case_rows"]
    write_q1_workbook(workbook_tables, output / "external_q1_tables.xlsx")
    write_artifact_manifest(output, config, f"EXTERNAL_{experiment}", seed)
    return results
